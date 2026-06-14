import shutil
from collections.abc import Generator
from pathlib import Path
from uuid import uuid4
from zipfile import ZipFile

import pytest
from docx import Document as DocxDocument

from app.services.document_text_extractor import DocumentTextExtractor


@pytest.fixture()
def parser_dir() -> Generator[Path, None, None]:
    path = Path("test_uploads") / f"parser-{uuid4()}"
    path.mkdir(parents=True, exist_ok=True)
    try:
        yield path
    finally:
        shutil.rmtree(path, ignore_errors=True)


def test_extract_docx_text_includes_paragraphs_and_tables(parser_dir: Path) -> None:
    file_path = parser_dir / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Contract paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Party A"
    table.cell(0, 1).text = "Party B"
    document.save(file_path)

    text = DocumentTextExtractor().extract_text(file_path)

    assert text is not None
    assert "Contract paragraph" in text
    assert "Party A\tParty B" in text


def test_extract_pdf_text(parser_dir: Path) -> None:
    file_path = parser_dir / "sample.pdf"
    file_path.write_bytes(_simple_pdf_bytes("Hello PDF contract"))

    text = DocumentTextExtractor().extract_text(file_path)

    assert text is not None
    assert "Hello PDF contract" in text


def test_extract_xlsx_text_includes_rows_and_shared_strings(parser_dir: Path) -> None:
    file_path = parser_dir / "sample.xlsx"
    _write_simple_xlsx(file_path)

    text = DocumentTextExtractor().extract_text(file_path)

    assert text is not None
    assert "Назва\tЗначення" in text
    assert "Площа\t120" in text


def test_extract_html_text_ignores_scripts(parser_dir: Path) -> None:
    file_path = parser_dir / "law.htm"
    file_path.write_text(
        """
        <html>
          <head><script>hidden()</script><style>.x { color: red; }</style></head>
          <body><h1>Конституція України</h1><p>Стаття 1. Україна є суверенна.</p></body>
        </html>
        """,
        encoding="utf-8",
    )

    text = DocumentTextExtractor().extract_text(file_path)

    assert text is not None
    assert "Конституція України" in text
    assert "Стаття 1" in text
    assert "hidden" not in text


def test_unsupported_file_returns_none(parser_dir: Path) -> None:
    file_path = parser_dir / "notes.txt"
    file_path.write_text("Plain text", encoding="utf-8")

    assert DocumentTextExtractor().extract_text(file_path) is None


def _simple_pdf_bytes(text: str) -> bytes:
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        (
            f"<< /Length {len(_pdf_text_stream(escaped_text))} >>\n"
            f"stream\n{_pdf_text_stream(escaped_text).decode('ascii')}\nendstream"
        ).encode("ascii"),
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, content in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(content)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


def _pdf_text_stream(text: str) -> bytes:
    return f"BT /F1 12 Tf 72 72 Td ({text}) Tj ET".encode("ascii")


def _write_simple_xlsx(file_path: Path) -> None:
    with ZipFile(file_path, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            """
            <Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
              <Default Extension="xml" ContentType="application/xml"/>
            </Types>
            """,
        )
        archive.writestr(
            "xl/sharedStrings.xml",
            """
            <sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
              <si><t>Назва</t></si>
              <si><t>Значення</t></si>
              <si><t>Площа</t></si>
            </sst>
            """,
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            """
            <worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
              <sheetData>
                <row r="1">
                  <c r="A1" t="s"><v>0</v></c>
                  <c r="B1" t="s"><v>1</v></c>
                </row>
                <row r="2">
                  <c r="A2" t="s"><v>2</v></c>
                  <c r="B2"><v>120</v></c>
                </row>
              </sheetData>
            </worksheet>
            """,
        )
