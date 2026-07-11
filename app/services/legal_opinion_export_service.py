import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sqlalchemy.orm import Session

from app.config import get_settings
from app.models.legal_opinion import LegalOpinion
from app.models.legal_opinion_export import LegalOpinionExport
from app.schemas.legal_opinion_schema import LegalOpinionExportResponse
from app.services.access_control import AccessControlService, WorkspacePermission

CONTENT_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


class LegalOpinionExportError(Exception):
    pass


class LegalOpinionExportService:
    def __init__(
        self,
        db: Session,
        access_control: AccessControlService | None = None,
    ) -> None:
        self.db = db
        self.access_control = access_control or AccessControlService(db)

    def export_opinion(
        self,
        *,
        opinion: LegalOpinion,
        user_id: str,
        export_format: str,
    ) -> LegalOpinionExportResponse:
        export_format = export_format.lower().strip()
        if export_format not in CONTENT_TYPES:
            raise LegalOpinionExportError("Unsupported export format")
        self.access_control.require_permission(
            workspace_id=opinion.workspace_id,
            user_id=user_id,
            permission=WorkspacePermission.READ,
        )
        output_dir = Path(get_settings().export_dir) / "legal_opinions"
        output_dir.mkdir(parents=True, exist_ok=True)
        base_name = self._safe_filename(f"legal-opinion-{opinion.id}")
        docx_path = output_dir / f"{base_name}.docx"
        self._write_docx(opinion, docx_path)
        final_path = docx_path
        if export_format == "pdf":
            final_path = self._convert_docx_to_pdf(docx_path, output_dir)
        file_size = final_path.stat().st_size
        export = LegalOpinionExport(
            legal_opinion_id=opinion.id,
            workspace_id=opinion.workspace_id,
            user_id=opinion.user_id,
            exported_by=user_id,
            export_format=export_format,
            file_path=str(final_path),
            content_type=CONTENT_TYPES[export_format],
            file_size=file_size,
            metadata_json={
                "review_status": opinion.review_status,
                "source_document_id": opinion.source_document_id,
            },
        )
        self.db.add(export)
        self.db.commit()
        self.db.refresh(export)
        return LegalOpinionExportResponse(
            ok=True,
            export_id=export.id,
            legal_opinion_id=opinion.id,
            export_format=export_format,
            file_path=export.file_path,
            content_type=export.content_type,
            file_size=export.file_size,
            message="Legal opinion exported.",
        )

    def _write_docx(self, opinion: LegalOpinion, path: Path) -> None:
        document = DocxDocument()
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        title = document.add_heading("Юридичний висновок", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"Статус рев'ю: {opinion.review_status}")
        document.add_paragraph(f"Workspace: {opinion.workspace_id}")
        document.add_paragraph(f"Opinion ID: {opinion.id}")
        document.add_heading("Запит", level=1)
        document.add_paragraph(opinion.question)
        document.add_heading("Відповідь", level=1)
        for block in opinion.answer.split("\n\n"):
            if block.strip():
                document.add_paragraph(block.strip())
        if opinion.sources_used:
            document.add_heading("Використані джерела", level=1)
            for key, value in opinion.sources_used.items():
                document.add_paragraph(f"{key}: {value}")
        if opinion.review_notes:
            document.add_heading("Нотатки рев'ю", level=1)
            document.add_paragraph(opinion.review_notes)
        document.save(path)

    def _convert_docx_to_pdf(self, docx_path: Path, output_dir: Path) -> Path:
        soffice = (
            shutil.which("soffice") or shutil.which("soffice.com") or shutil.which("libreoffice")
        )
        if not soffice:
            raise LegalOpinionExportError("LibreOffice/soffice is required for PDF export")
        with tempfile.TemporaryDirectory(prefix="jur-lo-") as profile_dir:
            subprocess.run(
                [
                    soffice,
                    f"-env:UserInstallation=file:///{Path(profile_dir).as_posix()}",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                text=True,
                timeout=120,
            )
        pdf_path = output_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.exists():
            raise LegalOpinionExportError("PDF conversion did not produce an output file")
        return pdf_path

    def _safe_filename(self, value: str) -> str:
        return re.sub(r"[^A-Za-z0-9_.-]+", "-", value).strip("-.") or "legal-opinion"
