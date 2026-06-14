from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document as DocxDocument
from pypdf import PdfReader

PDF_EXTENSIONS = frozenset({".pdf"})
DOCX_EXTENSIONS = frozenset({".docx"})
XLSX_EXTENSIONS = frozenset({".xlsx"})
HTML_EXTENSIONS = frozenset({".html", ".htm"})
PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
HTML_MIME_TYPES = frozenset({"text/html", "application/xhtml+xml"})


class UnsupportedDocumentTypeError(Exception):
    pass


class DocumentTextExtractionError(Exception):
    pass


class DocumentTextExtractor:
    def extract_text(self, file_path: str | Path, content_type: str | None = None) -> str | None:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix in PDF_EXTENSIONS or content_type == PDF_MIME_TYPE:
            return self.extract_pdf_text(path)
        if suffix in DOCX_EXTENSIONS or content_type == DOCX_MIME_TYPE:
            return self.extract_docx_text(path)
        if suffix in XLSX_EXTENSIONS or content_type == XLSX_MIME_TYPE:
            return self.extract_xlsx_text(path)
        if suffix in HTML_EXTENSIONS or content_type in HTML_MIME_TYPES:
            return self.extract_html_text(path)
        return None

    def extract_pdf_text(self, file_path: str | Path) -> str:
        try:
            reader = PdfReader(str(file_path))
            pages = [page.extract_text() or "" for page in reader.pages]
        except Exception as exc:
            raise DocumentTextExtractionError("Could not extract text from PDF") from exc
        return _normalize_text("\n".join(pages))

    def extract_docx_text(self, file_path: str | Path) -> str:
        try:
            document = DocxDocument(str(file_path))
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
            for table in document.tables:
                for row in table.rows:
                    row_text = "\t".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
        except Exception as exc:
            raise DocumentTextExtractionError("Could not extract text from DOCX") from exc
        return _normalize_text("\n".join(parts))

    def extract_xlsx_text(self, file_path: str | Path) -> str:
        try:
            with ZipFile(file_path) as archive:
                shared_strings = _read_shared_strings(archive)
                worksheet_names = sorted(
                    name
                    for name in archive.namelist()
                    if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
                )
                rows: list[str] = []
                for worksheet_name in worksheet_names:
                    rows.extend(_read_worksheet_rows(archive, worksheet_name, shared_strings))
        except Exception as exc:
            raise DocumentTextExtractionError("Could not extract text from XLSX") from exc
        return _normalize_text("\n".join(rows))

    def extract_html_text(self, file_path: str | Path) -> str:
        try:
            raw = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            parser = _HTMLTextParser()
            parser.feed(raw)
        except Exception as exc:
            raise DocumentTextExtractionError("Could not extract text from HTML") from exc
        return _normalize_text(parser.text)


class _HTMLTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._ignored_depth = 0

    @property
    def text(self) -> str:
        return "\n".join(self.parts)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._ignored_depth += 1
        if tag in {"br", "p", "tr", "li", "div", "section", "article", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._ignored_depth:
            self._ignored_depth -= 1
        if tag in {"p", "tr", "li", "div", "section", "article", "h1", "h2", "h3"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._ignored_depth:
            return
        text = data.strip()
        if text:
            self.parts.append(text)


def _normalize_text(value: str) -> str:
    lines = [line.strip() for line in value.replace("\r\n", "\n").split("\n")]
    return "\n".join(line for line in lines if line)


def _read_shared_strings(archive: ZipFile) -> list[str]:
    try:
        raw = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    root = ElementTree.fromstring(raw)
    return [_cell_text(item) for item in root]


def _read_worksheet_rows(
    archive: ZipFile, worksheet_name: str, shared_strings: list[str]
) -> list[str]:
    root = ElementTree.fromstring(archive.read(worksheet_name))
    rows: list[str] = []
    for row in root.iter(_tag("row")):
        values = [
            value
            for cell in row.iter(_tag("c"))
            if (value := _read_cell_value(cell, shared_strings))
        ]
        if values:
            rows.append("\t".join(values))
    return rows


def _read_cell_value(cell: ElementTree.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t")
    value_node = cell.find(_tag("v"))
    inline_node = cell.find(_tag("is"))
    if cell_type == "inlineStr" and inline_node is not None:
        return _cell_text(inline_node)
    if value_node is None or value_node.text is None:
        return ""
    raw_value = value_node.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(raw_value)]
        except (IndexError, ValueError):
            return ""
    return raw_value


def _cell_text(element: ElementTree.Element) -> str:
    return "".join(text.strip() for text in element.itertext() if text.strip())


def _tag(name: str) -> str:
    return f"{{http://schemas.openxmlformats.org/spreadsheetml/2006/main}}{name}"
